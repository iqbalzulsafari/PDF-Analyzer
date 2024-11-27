import streamlit as st
from jamaibase import JamAI, protocol as p
import os
from docx import Document
from io import BytesIO
import random
import string
from PyPDF2 import PdfReader
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Fetch token and project ID from environment variables
ACCESS_TOKEN = os.getenv("ACCESS_TOKEN")
PROJECT_ID = os.getenv("PROJECT_ID")

jamai = JamAI(token=ACCESS_TOKEN, project_id=PROJECT_ID)

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


# Function to generate a random filename
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"


# Set up the Streamlit app
st.set_page_config(page_title="PDF Analyzer", page_icon="📝")
st.title("🌟 PDF Analyzer - Your AI Assistant for summarization")

# Custom CSS to style the UI
st.markdown(
    """
    <style>
    body {
        background-color: #1e1e1e;
        color: #f0f0f0;
    }
    .generated-output {
        background-color: #444;
        padding: 15px;
        border-radius: 10px;
        margin-top: 20px;
        box-shadow: 0px 4px 10px rgba(0, 0, 0, 0.5);
        color: #f0f0f0;
    }
    .generated-output h4 {
        color: #FFA500;
    }
    .stButton>button {
        background-color: #4CAF50;
        color: white;
        border: none;
        border-radius: 5px;
        padding: 10px 20px;
        cursor: pointer;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Containers for inputs
with st.container():
    st.header("📄 Upload the PDF file")
    # Upload PDF CV
    pdf = st.file_uploader("Upload the pdf file", type="pdf")

# Action to process inputs
if st.button("🚀 Process Input", use_container_width=True):
    if pdf:
        # Extract text from CV PDF
        pdf_text = extract_text_from_pdf(pdf)

        # Add rows to the existing table with the input data
        try:
            completion = jamai.add_table_rows(
                "action",
                p.RowAddRequest(
                    table_id="pdf-analysis",
                    data=[{"PDF_file": pdf_text}],
                    stream=False
                )
            )

            # Display the output generated in the columns
            if completion.rows:
                output_row = completion.rows[0].columns
                summary = output_row.get("summary")

                st.subheader("✨ Generated Output")
                st.markdown(
                    f"""
                    <div class="generated-output">
                        <h4>📝 Summary:</h4> <p>{summary.text if summary else 'N/A'}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Download the final report as a .docx file
                with st.container():
                    st.subheader("📥 Download Final Report")
                    doc = Document()
                    doc.add_heading("File Report", level=1)
                    # Summary and Work Experience Section
                    doc.add_heading("Summary", level=2)
                    doc.add_paragraph(summary.text if summary else 'N/A')

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="📄 Download Final Report as .docx",
                        data=buffer,
                        file_name=generate_random_filename(),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("⚠️ Failed to get a response. Please try again.")
        except Exception as e:
            st.error(f"❌ An error occurred: {e}")
    else:
        st.warning("⚠️ Please upload a PDF file.")